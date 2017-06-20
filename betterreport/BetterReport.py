"""
Name:	BetterReport
Author:	John Rachwan

Basic Command: py GoReport.py --id <Campaign id> --format csv
"""

# Basic imports
from gophish import Gophish
import click
import sys
import csv
import configparser
import time
import random
# Imports for statistics, e.g. browsera and operating systems
from user_agents import parse
from collections import Counter
# Imports for web requests, e.g. Google Maps API for location data
# Disables the insecure HTTPS warning for the self-signed GoPhish certs
import requests
from requests.packages.urllib3.exceptions import InsecureRequestWarning

requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
# Import the MaxmInd's GeoLite for IP address GeoIP look-ups
from geolite2 import geolite2
# Imports for writing the Word.doc report
import os.path
from docx import *

from docx import Document
from docx.shared import Inches

document = Document()

from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import matplotlib.pyplot as plt

# Open the config file to make sure it exists and is readable
try:
	config = configparser.ConfigParser()
	config.read('gophish.ini')
except Exception as e:
	print("[!] Could not open the /gophish.ini config file -- make sure it exists and is readable.")
	print("L.. Details: {}".format(e))
	sys.exit()


def ConfigSectionMap(section):
	"""This function helps by reading accepting a config file section, from gophish.ini,
	and returning a dictionary object that can be referenced for configuration settings.
	"""
	section_dict = {}
	options = config.options(section)
	for option in options:
		try:
			section_dict[option] = config.get(section, option)
			if section_dict[option] == -1:
				DebugPrint("[-] Skipping: {}".format(option))
		except:
			print("[!] There was an error with: {}".format(option))
			section_dict[option] = None
	return section_dict


# Read in the values from the config file
try:
	GP_HOST = ConfigSectionMap("GoPhish")['gp_host']
	API_KEY = ConfigSectionMap("GoPhish")['api_key']
	# TODO: Allow specifying an MMDB file location
	# MMDB = ConfigSectionMap("GeoIP")['mmdb_path']
except Exception as e:
	print("[!] There was a problem reading values from the gophish.ini file!")
	print("L.. Details: {}".format(e))
	sys.exit()


def set_column_width(column, width):
	"""Custom function for quickly and easily setting the width of a table's
	column in the Word docx output.

	"""
	for cell in column.cells:
		cell.width = width

# Setup an AliasedGroup for CLICK
class AliasedGroup(click.Group):
	"""Allows commands to be called by their first unique character."""

	def get_command(self, ctx, cmd_name):
		"""
		Allows commands to be called by thier first unique character
		:param ctx: Context information from click
		:param cmd_name: Calling command name
		:return:
		"""
		rv = click.Group.get_command(self, ctx, cmd_name)
		if rv is not None:
			return rv
		matches = [x for x in self.list_commands(ctx)
			if x.startswith(cmd_name)]
		if not matches:
			return None
		elif len(matches) == 1:
			return click.Group.get_command(self, ctx, matches[0])
		ctx.fail('Too many matches: %s' % ', '.join(sorted(matches)))


# Create the help option for CLICK
CONTEXT_SETTINGS = dict(help_option_names=['-h', '--help'])
@click.group(cls=AliasedGroup, context_settings=CONTEXT_SETTINGS)


def GoReport():
	# Everything starts here
	pass


# Setup our CLICK arguments and help text
@GoReport.command(name='report', short_help="Generate a full report for the selected campaign -- either CSV or DOCX.")
@click.option('--id', help="The target campaign's ID.", required=False, multiple=True)
@click.option('--format', type=click.Choice(['csv', 'word', 'quick']), help="Use this option to choose between report formats.", required=False)
@click.pass_context
def parse_options(self, id, format):
	"""GoReport uses the GoPhish API to connect to your GoPhish instance using the
	IP address, port, and API key for your installation. This information is provided
	in the gophish.ini file and loaded at runtime. GoReport will collect details
	for the specified campaign and output statistics and interesting data for you.

	Target a campaign by its ID number with --id and then select a report format.\n
	   * csv: A comma separated file. Good for copy/pasting into other documents.\n
	   * word: A formatted docx file. A template.docx file is required (see the README).\n
	   * quick: Command line output of some basic stats. Good for a quick check or client call.\n

	Hint: Multiple reports can be run by adding additional --id arguments!
	"""

	# Proceed with one campaign ID at a time
	id = raw_input('Enter a id: ')
	format = raw_input('Enter a group name: ')
	for x in id:
		# CAM_ID is double declared like this so there are no complaints if this check fails
		CAM_ID = x
		try:
			# Test to make sure the provided ID is really an integer
			CAM_ID = int(CAM_ID)
		except:
			print("[!] You entered an invalid campaign ID! {} will not do!".format(CAM_ID))
			sys.exit()

		# Kick-off a new campaign object with our options
		try:
			gophish = GPCampaign(CAM_ID, format)
			gophish.run()
		except:
			# Oops? Hopefully the exception was caught elsewhere and we can continue.
			continue


# Everything from here on is the GoPhish Campaign class
class GPCampaign(object):
	"""This class uses the GoPhish library to create a new GoPhish API connection
	and queries GoPhish for information and results related to the specified
	campaign.
	"""
	# Variables for holding GoPhish models
	campaign = None
	results = None
	timeline = None

	# Variables for holding campaign information
	cam_id = None
	cam_name = None
	cam_status = None
	created_date = None
	launch_date = None
	completed_date = None
	cam_url = None
	cam_redirect_url = None
	cam_from_address = None
	cam_subject_line = None
	cam_template_name = None
	cam_capturing_passwords = None
	cam_capturing_credentials = None
	cam_page_name = None
	cam_smtp_host = None

	# Variables and lists for tracking event numbers
	total_targets = None
	total_sent = None
	total_opened = None
	total_clicked = None
	total_submitted = None
	targets_opened = []
	targets_clicked = []
	targets_submitted = []

	# Lists for holding totals for statistics
	browsers = []
	operating_systems = []
	locations = []
	ip_addresses = []

	# Output options
	OUTPUT_TYPE = None
	output_csv_report = None
	output_word_report = None

	def __init__(self, CAM_ID, OUTPUT_TYPE):
		"""Initiate the connection to the GoPhish server with the provided host, port, and API key"""
		self.OUTPUT_TYPE = OUTPUT_TYPE
		try:
			# Connect to the GoPhish API
			# NOTE: This step succeeds even with a bad API key, so the true test is fetching an ID
			print("[+] Connecting to GoPhish at {}".format(GP_HOST))
			self.api = Gophish(API_KEY, host=GP_HOST, verify=False)
			# Request the details for the provided campaign ID
			print("[+] We will now try fetching results for Campaign ID {}.".format(CAM_ID))
			self.campaign = self.api.campaigns.get(campaign_id=CAM_ID)
			# Check to see is a success message was returned with a message
			# This will mean there is probably a problem with the connection or API key
			try:
				if self.campaign.success is False:
					print("[!] Connection to GoPhish failed!")
					print("L.. Details: {}".format(self.campaign.message))
			# If self.campaign.success does not exist then we were successful
			except:
				print("[+] We have successfully pulled the campaign details for ID {}.".format(CAM_ID))
		except Exception as e:
			if self.campaign.success:
				print("yay?")
			print("[!] There was a problem fetching this campaign ID's details. Are you sure your URL and API key are correct? Check HTTP vs HTTPS!".format(CAM_ID))
			print("L.. Details: {}".format(e))
			sys.exit()

		# Create the MaxMind GeoIP reader for the CeoLite2-City.mmdb database file
		self.geoip_reader = geolite2.reader()

	def run(self):
		"""Run everything to process the target campaign."""
		# Collect campaign details and process data
		self.collect_campaign_details()
		self.parse_results()
		self.parse_timeline_events()
		# Generate the report
		if self.OUTPUT_TYPE == "csv":
			print("[+] Building the report -- you selected a csv report.")
			self.output_csv_report = self._build_output_csv_file_name()
			self.write_csv_report()
		elif self.OUTPUT_TYPE == "word":
			print("[+] Building the report -- you selected a Word/docx report.")
			print("[+] Looking for the template.docx to be used for the Word report.")
			if os.path.isfile("template.docx"):
				self.output_word_report = self._build_output_word_file_name()
				self.write_word_report()
			else:
				print("[!] Could not find the template document! Make sure 'template.docx' is in the GoReport directory.")
				sys.exit()
		elif self.OUTPUT_TYPE == "quick":
			print("[+] No report this time. Here are your quick stats:\n")
			self.get_quick_stats()

	def get_quick_stats(self):
		"""Present quick stats for the campaign. Just basic numbers and some details."""
		print(self.campaign.name)
		print("Status:\t\t{}".format(self.cam_status))
		print("Created:\t{} on {}".format(self.created_date.split("T")[1].split(".")[0], self.created_date.split("T")[0]))
		print("Started:\t{} on {}".format(self.launch_date.split("T")[1].split(".")[0], self.launch_date.split("T")[0]))
		if self.cam_status == "Completed":
			print("Completed:\t{} on {}".format(self.completed_date.split("T")[1].split(".")[0], self.completed_date.split("T")[0]))
		print("Total targets:\t{}".format(self.total_targets))
		print("Emails sent:\t{}".format(self.total_sent))
		print("Opened events:\t{}".format(self.total_opened))
		print("Click events:\t{}".format(self.total_clicked))
		print("Entered data:\t{}".format(self.total_submitted))
		print("IPs seen:\t{}".format(len(self.ip_addresses)))

	def _build_output_csv_file_name(self):
		"""A helper function to create the output report name."""
		csv_report = "GoPhish Results for Campaign - {}.csv".format(self.cam_name)
		return csv_report

	def _build_output_word_file_name(self):
		"""A helper function to create the output report name."""
		word_report = "GoPhish Results for Campaign - {}.docx".format(self.cam_name)
		return word_report

	def collect_campaign_details(self):
		"""Collect the campaign's details set values for each of the declared variables."""
		# Collect the basic campaign details
		# Plus a quick and dirty check to see if the campaign ID is valid
		try:
			self.cam_id = self.campaign.id
		except:
			print("[!] Looks like that campaign ID does not exist! Skipping it...")

		self.cam_name = self.campaign.name
		self.cam_status = self.campaign.status
		self.created_date = self.campaign.created_date
		self.launch_date = self.campaign.launch_date
		self.completed_date = self.campaign.completed_date
		self.cam_url = self.campaign.url

		# Collect the results and timeline, lists
		self.results = self.campaign.results
		self.timeline = self.campaign.timeline

		# Collect SMTP information
		self.smtp = self.campaign.smtp
		self.cam_from_address = self.smtp.from_address
		self.cam_smtp_host = self.smtp.host

		# Collect the template information
		self.template = self.campaign.template
		self.cam_subject_line = self.template.subject
		self.cam_template_name = self.template.name
		self.cam_template_attachments = self.template.attachments

		# Collect the landing page information
		self.page = self.campaign.page
		self.cam_page_name = self.page.name
		self.cam_redirect_url = self.page.redirect_url
		self.cam_capturing_passwords = self.page.capture_passwords
		self.cam_capturing_credentials = self.page.capture_credentials

	def parse_results(self):
		"""Process the results model to collect basic data, like total targets.

		The results model can provide:
		first_name, last_name, email, position, and IP address
		"""
		# Total length of results gives us the total number of targets
		self.total_targets = len(self.results)

		# Go through all results and extract data for statistics
		for x in self.results:
			if not x.ip == "":
				self.ip_addresses.append(x.ip)

	def parse_timeline_events(self):
		"""Process the timeline model to collect basic data, like total clicks.

		The timeline model contains all events that occured during the campaign.
		"""
		# Create counters for enumeration
		sent_counter = 0
		opened_counter = 0
		click_counter = 0
		submitted_counter = 0
		# Run through all events and count each of the four basic events
		for x in self.timeline:
			if x.message == "Email Sent":
				sent_counter += 1
			elif x.message == "Email Opened":
				opened_counter += 1
				self.targets_opened.append(x.email)
			elif x.message == "Clicked Link":
				click_counter += 1
				self.targets_clicked.append(x.email)
			elif x.message == "Submitted Data":
				submitted_counter += 1
				self.targets_submitted.append(x.email)
		# Assign the counter values to our tracking lists
		self.total_sent = sent_counter
		self.total_opened = opened_counter
		self.total_clicked = click_counter
		self.total_submitted = submitted_counter

	def lookup_ip(self, ip):
		"""Check the GeoLite database for a location for the provided IP address.

		This returns a large dict with more data than is probably needed for
		a report. This gets continent, country, registered_country, and location.
		Also, this dict includes multiple languages.

		You may wonder why get_google_location_data() is needed if this provides
		a lot of data from MaxMind. Unfortunately, the MaxMind database will not
		always have the data needed most for the report (city, state, country).
		It may only have the continent name. Luckily, it seems to always have coordinates
		that can be compared to GoPhish's coordinates and passed to get_google_location_data().
		"""
		match  = self.geoip_reader.get(ip)
		if match is not None:
			return match
		else:
			# return "No match"
			return None

	def get_google_location_data(self, lat, lon):
		"""Use Google's Maps API to collect GeoIP info for the provided latitude
		and longitude.

		Google returns a bunch of JSON with a variety of location data.
		This function sticks to the first set of "address_components" for the
		country, locality (city), and administrative_level_1 (state).

		Ex: http://maps.googleapis.com/maps/api/geocode/json?latlng=35,-93&sensor=false
		"""
		url = "http://maps.googleapis.com/maps/api/geocode/json?latlng={},{}&sensor=false".format(lat, lon)
		v = requests.get(url)
		j = v.json()
		try:
			# Get the first set of 'address_components' from the JSON results
			components = j['results'][0]['address_components']
			country = town = None
			for c in components:
				if "country" in c['types']:
					country = c['long_name']
				if "locality" in c['types']:
					town = c['long_name']
				if "administrative_area_level_1" in c['types']:
					state = c['long_name']
			# TODO: Remove the periods and allow for commas for the Word output
			# The commas are a problem for csv output, but can be replaced for that report format
			return "{}. {}. {}".format(town, state, country)
		except:
			# return "None"
			return None

	def compare_ip_addresses(self, target_ip, browser_ip):
		"""Compare the IP addresses of the target to that of an event. The goal:
		Looking for a mismatch that might identify some sort of interesting event.
		This might indicate an email was forwarded, a VPN was switched on/off, or
		maybe the target is at home.
		"""
		if target_ip == browser_ip:
			return target_ip
		else:
			# We have an IP mismatch! Hard to tell what this might be.
			print("[!] Interesting Event: The target's IP does not match their browser's IP!")
			print("L.. This target's ({}) URL was clicked from a browser at {} -- email may have been forwarded or the target is home/using VPN/etc. Interesting!".format(target_ip, browser_ip))
			# This is an IP address not included in the results model, so we add it to our list here
			self.ip_addresses.append(browser_ip)
			return browser_ip

	def compare_ip_coordinates(self, target_latitude, target_longitude, mmdb_latitude, mmdb_longitude, ip_address):
		"""Compare the IP address cooridnates reported by MaxMind and GoPhish.
		If they do not match, some additional -- manual -- investigation should
		be done for any client-facing deliverables.
		"""
		if target_latitude == mmdb_latitude and target_longitude == mmdb_longitude:
			# Coordinates match what GoPhish recorded, so query Google Maps for details
			coordinates_location = self.get_google_location_data(target_latitude, target_longitude)
			if not coordinates_location is None:
				self.locations.append(coordinates_location)
				return coordinates_location
			else:
				return "Google timeout"
		else:
			# MaxMind and GoPhish have different coordinates, so this is a tough spot
			# Both locations can be recorded for investigation, but what to do for location statistics?
			# It was decided both would be recorded as one location with an asterisk, flagged for investigation
			print("[!] Warning: Location coordinates mis-match between MaxMind and GoPhish for {}. Look for location with * to investigate and pick the right location.".format(ip_address))
			coordinates_location = self.get_google_location_data(target_latitude, target_longitude)
			# Sleep between checks to avoid bad Results
			# In cases with a lot of mismatches Google seems to return no reuslts for back-to-back requests
			time.sleep(2)
			alt_coordinates_location = self.get_google_location_data(mmdb_latitude, mmdb_longitude)
			if not alt_coordinates_location is None and not coordinates_location is None:
				coordinates_location += "     ALTERNATE:{}".format(alt_coordinates_location)
			elif not coordinates_location is None and alt_coordinates_location is None:
				coordinates_location += "     ALTERNATE: MaxMind returned No Results"
			elif coordinates_location is None and not alt_coordinates_location is None:
				coordinates_location = alt_coordinates_location

			self.locations.append(coordinates_location + " *")
			return "{}".format(coordinates_location + " *")

	def write_csv_report(self):
		"""Assemble and output the csv file report.

		Throughout this function, results are assembled by adding commas and Then
		adding to the results string, i.e. 'result_A' and then 'result_A' += ',result_B'.
		This is so the result can be written to the csv file and have the different
		results end up in the correct columns.
		"""
		with open(self.output_csv_report, 'w') as csvfile:
			# Create the csv writer
			writer = csv.writer(csvfile, dialect='excel', delimiter=',', quotechar="'", quoting=csv.QUOTE_MINIMAL)

			# Write a campaign summary at the top of the report
			writer.writerow(["CAMPAIGN RESULTS FOR:", "{}".format(self.cam_name)])
			writer.writerow(["Status", "{}".format(self.cam_status)])
			writer.writerow(["Created", "{}".format(self.created_date)])
			writer.writerow(["Started", "{}".format(self.launch_date)])
			# If the campaign has been completed, we will record that, too
			if self.cam_status == "Completed":
				writer.writerow(["Completed", "{}".format(self.completed_date)])
			# Write the campaign details -- email details and template settings
			writer.writerow("")
			writer.writerow(["CAMPAIGN DETAILS"])
			writer.writerow(["From", "{}".format(self.cam_from_address)])
			writer.writerow(["Subject", "{}".format(self.cam_subject_line)])
			writer.writerow(["Phish URL", "{}".format(self.cam_url)])
			if self.cam_redirect_url == "":
				writer.writerow(["Redirect URL", "Not Used"])
			else:
				writer.writerow(["Redirect URL", "{}".format(self.cam_redirect_url)])
			if self.cam_template_attachments == []:
				writer.writerow(["Attachment(s)", "None"])
			else:
				writer.writerow(["Attachment(s)", "{}".format(self.cam_template_attachments)])
			writer.writerow(["Captured Credentials", "{}".format(self.cam_capturing_credentials)])
			writer.writerow(["Stored Passwords", "{}".format(self.cam_capturing_passwords)])
			# Write a high level summary for stats
			writer.writerow("")
			writer.writerow(["HIGH LEVEL RESULTS"])
			writer.writerow(["Total Targets", "{}".format(self.total_targets)])
			writer.writerow(["Opened", "{}".format(self.total_opened)])
			writer.writerow(["Clicked", "{}".format(self.total_clicked)])
			writer.writerow(["Entered Data", "{}".format(self.total_submitted)])



			# End of the campaign summary and beginning of the event summary
			writer.writerow("")
			writer.writerow(["SUMMARY OF EVENTS"])
			writer.writerow(["Email Address", "Open", "Click", "Phish"])
			# Add targets to the results table
			for target in self.results:
				result = target.email
				# Chck if this target was recorded as viewing the email (tracking image)
				if target.email in self.targets_opened:
					result += ",Y"
				else:
					result += ",N"
				# Check if this target clicked the link
				if target.email in self.targets_clicked:
					result += ",Y"
				else:
					result += ",N"
				# Check if this target submitted data
				if target.email in self.targets_submitted:
					result += ",Y"
				else:
					result += ",N"

				writer.writerow(["{}".format(result)])

			# End of the event summary and beginning of the detailed results
			for target in self.results:
				writer.writerow("")
				writer.writerow(["{} {}".format(target.first_name, target.last_name, target.email)])
				writer.writerow(["{}".format(target.email)])
				# Parse each timeline event
				# Timestamps are parsed to get date and times by splitting date
				# and time and dropping the milliseconds and timezone
				# Ex: 2017-01-30T14:31:22.534880731-05:00
				for event in self.timeline:
					if event.message == "Email Sent" and event.email == target.email:
						# Parse the timestamp into separate date and time variables
						temp = event.time.split('T')
						sent_date = temp[0]
						sent_time = temp[1].split('.')[0]
						# Record the email sent date and time in the report
						writer.writerow(["Sent on {} at {}".format(sent_date.replace(",", ""), sent_time)])

					if event.message == "Email Opened" and event.email == target.email:
						# Parse the timestamp
						temp = event.time.split('T')
						# Record the email preview date and time in the report
						writer.writerow(["Email Preview",  "{} {}".format(temp[0], temp[1].split('.')[0])])

					if event.message == "Clicked Link" and event.email == target.email:
						# Parse the timestmap and add the time to the results row
						temp = event.time.split('T')
						result = temp[0] + " " + temp[1].split('.')[0]

						# Add the IP address to the results row
						# Sanity check to see if browser IP matches the target's recorded IP
						result += ",{}".format(self.compare_ip_addresses(target.ip, event.details['browser']['address']))

						# Get the location data and add to results row
						# This is based on the IP address pulled from the browser for this event
						# Start by getting the coordinates from GeoLite2
						mmdb_location = self.lookup_ip(event.details['browser']['address'])
						if not mmdb_location == None:
							mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'], mmdb_location['location']['longitude']
							# Check if GoPhish's coordinates agree with these MMDB results
							result += ",{}".format(self.compare_ip_coordinates(target.latitude, target.longitude, mmdb_latitude, mmdb_longitude, event.details['browser']['address']))
						else:
							result += ",IP address look-up returned None"

						# Parse the user-agent string and add browser and OS details to the results row
						user_agent = parse(event.details['browser']['user-agent'])

						browser_details = user_agent.browser.family + " " + user_agent.browser.version_string
						result += ",{}".format(browser_details)
						self.browsers.append(browser_details)

						os_details = user_agent.os.family + " " + user_agent.os.version_string
						result += ",{}".format(os_details)
						self.operating_systems.append(os_details)

						# Write the results row to the report for this target
						writer.writerow(["Email Link Clicked"])
						writer.writerow(["Time", "IP", "City", "Browser", "Operating System"])
						writer.writerow([result])

					# Now we have events for submitted data. A few notes on this:
					# There is no expectation of data being submitted without a Clicked Link event
					# Assuming that, the following process does NOT flag IP
					# mismatches or add to the list of seen locations, OSs, IPs, or browsers.
					if event.message == "Submitted Data" and event.email == target.email:
						# Parse the timestmap and add the time to the results row
						temp = event.time.split('T')
						result = temp[0] + " " + temp[1].split('.')[0]

						# Add the IP address to the results row
						result += ",{}".format(event.details['browser']['address'])

						# Get the location data and add to results row
						# This is based on the IP address pulled from the browser for this event
						# Start by getting the coordinates from GeoLite2
						mmdb_location = self.lookup_ip(event.details['browser']['address'])
						if not mmdb_location == None:
							mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'], mmdb_location['location']['longitude']
							# Check if GoPhish's coordinates agree with these MMDB results
							loc = self.compare_ip_coordinates(target.latitude, target.longitude, mmdb_latitude, mmdb_longitude, event.details['browser']['address'])
							if not loc is None:
								result += loc
							else:
								result += "None"
						else:
							result += ",IP address look-up returned None"

						# Parse the user-agent string and add browser and OS details to the results row
						user_agent = parse(event.details['browser']['user-agent'])

						browser_details = user_agent.browser.family + " " + user_agent.browser.version_string
						result += ",{}".format(browser_details)

						os_details = user_agent.os.family + " " + user_agent.os.version_string
						result += ",{}".format(os_details)

						# Get just the submitted data from the event's payload
						submitted_data = ""
						data_payload = event.details['payload']
						# Get all of the submitted data
						for key, value in data_payload.items():
							# To get just submitted data, we drop the 'rid' key
							if not key == "rid":
								submitted_data += "{}:{}   ".format(key, str(value).strip("[").strip("]"))

						result += ",{}".format(submitted_data)
						# Write the results row to the report for this target
						writer.writerow(["Submitted Data Captured"])
						writer.writerow(["Time", "IP", "City", "Browser", "Operating System", "Data Captured"])
						writer.writerow([result])

			# End of the detailed results and the beginning of browser, location, and OS stats
			# Counter is used to count all elements in the lists to create a unique list with totals
			writer.writerow("")
			writer.writerow(["RECORDED BROWSERS BY UA:"])
			writer.writerow(["Browser", "Seen"])

			counted_browsers = Counter(self.browsers)
			for key, value in counted_browsers.items():
				writer.writerow(["{},{}".format(key, value)])

			writer.writerow("")
			writer.writerow(["RECORDED OP SYSTEMS:"])
			writer.writerow(["Operating System", "Seen"])

			counted_os = Counter(self.operating_systems)
			for key, value in counted_os.items():
				writer.writerow(["{},{}".format(key, value)])

			writer.writerow([" "])
			writer.writerow(["RECORDED LOCATIONS:"])
			writer.writerow(["Location", "Visits"])

			counted_locations = Counter(self.locations)
			for key, value in counted_locations.items():
				writer.writerow(["{},{}".format(key, value)])

			writer.writerow([" "])
			writer.writerow(["RECORDED IP ADDRESSES:"])
			writer.writerow(["IP Address", "Seen"])

			counted_ip_addresses = Counter(self.ip_addresses)
			for key, value in counted_ip_addresses.items():
				writer.writerow(["{},{}".format(key, value)])

			print("[+] Done! Check \'{}\' for your results.".format(self.output_csv_report))



	def write_word_report(self):
		"""Assemble and output the csv file report."""
		# Create document writer using the template and a style editor
		d = Document("template.docx")
		styles = d.styles

		# Create a custom style for table cells
		style = styles.add_style('Cell Text', WD_STYLE_TYPE.CHARACTER)
		cellText = d.styles['Cell Text']
		cellText_font = cellText.font
		cellText_font.name = 'Calibri'
		cellText_font.size = Pt(12)
		cellText_font.bold = True
		cellText_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

		# Write a campaign summary at the top of the report
		d.add_heading("Executive Summary", 1)
		p = d.add_paragraph()
		run = p.add_run("CAMPAIGN RESULTS FOR: {}".format(self.cam_name))
		run.bold = True
		# Runs are basically "runs" os text and must be aligned in the completed_date
		# like we want them aligned in the report -- thus they are pushed left
		p.add_run("""
Status: {}
Created: {}
Started: {}
Completed: {}
""".format(self.cam_status, self.created_date, self.launch_date,
		self.completed_date))

		# Write the campaign details -- email details and template settings
		run = p.add_run("CAMPAIGN DETAILS")
		run.bold = True

		p.add_run("""
From: {}
Subject: {}
Phish URL: {}
Redirect URL: {}
Attachment(s): {}
Captured Credentials: {}
Stored Passwords: {}
""".format(self.cam_from_address, self.cam_subject_line, self.cam_url,
		self.cam_redirect_url, self.cam_template_attachments, self.cam_capturing_credentials,
		self.cam_capturing_passwords))

		# Write a high level summary for stats
		run = p.add_run("HIGH LEVEL RESULTS")
		run.bold = True

		p.add_run("""
Total Targets: {}
Opened: {}
Clicked: {}
Entered Data: {}
""".format(self.total_targets, self.total_opened, self.total_clicked,
		self.total_submitted))
		##########################################################################
		clicked = (self.total_clicked / self.total_targets) * 100
		not_clicked = 100 - clicked
		labels = 'Not Clicked', 'Clicked'
		sizes = [not_clicked, clicked]
		explode = (0, 0.1)  # only "explode" the 2nd slice (i.e. 'Hogs')

		fig1, ax1 = plt.subplots()
		ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%',
				shadow=True, startangle=90)
		ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

		# plt.show()
		plt.savefig("john", dpi=150, facecolor='w', edgecolor='w',
					orientation='portrait', papertype=None, format=None,
					transparent=True, bbox_inches=None, pad_inches=0.1,
					frameon=None)
		d.add_picture('john.png', width=Inches(5))
		##################################################
		d.add_page_break()

		# End of the campaign summary and beginning of the event summary
		d.add_heading("Summary of Events", 1)
		d.add_paragraph("The table below summarizes who opened and clicked on emails sent in this campaign.")

		# Create a table to hold the event summary results
		table = d.add_table(rows=1, cols=4, style="GoReport")
		set_column_width(table.columns[0], Cm(4.2))
		set_column_width(table.columns[1], Cm(1.4))
		set_column_width(table.columns[2], Cm(1.4))
		set_column_width(table.columns[3], Cm(1.4))

		header1 = table.cell(0,0)
		header1.text = ""
		header1.paragraphs[0].add_run("Email Address", "Cell Text").bold = True

		header2 = table.cell(0,1)
		header2.text = ""
		header2.paragraphs[0].add_run("Open", "Cell Text").bold = True

		header3 = table.cell(0,2)
		header3.text = ""
		header3.paragraphs[0].add_run("Click", "Cell Text").bold = True

		header4 = table.cell(0,3)
		header4.text = ""
		header4.paragraphs[0].add_run("Phish", "Cell Text").bold = True

		# Add targets to the results table
		counter = 1
		for target in self.results:
			table.add_row()
			email_cell = table.cell(counter,0)
			email_cell.text = "{}".format(target.email)

			if target.email in self.targets_opened:
				temp_cell = table.cell(counter,1)
				temp_cell.text = "Y"
			else:
				temp_cell = table.cell(counter,1)
				temp_cell.text = "N"

			if target.email in self.targets_clicked:
				temp_cell = table.cell(counter,2)
				temp_cell.text = "Y"
			else:
				temp_cell = table.cell(counter,2)
				temp_cell.text = "N"

			if target.email in self.targets_submitted:
				temp_cell = table.cell(counter,3)
				temp_cell.text = "Y"
			else:
				temp_cell = table.cell(counter,3)
				temp_cell.text = "N"

			counter += 1

		d.add_page_break()

		# End of the event summary and beginning of the detailed results
		d.add_heading("Detailed Findings", 1)
		for target in self.results:
			# Create counters that will be used when adding rows
			# We need a counter to track the cell location
			opened_counter = 1
			clicked_counter = 1
			submitted_counter = 1
			# Create a heading 1 for the first and last name and heading 2 for email address
			d.add_heading("{} {}".format(target.first_name, target.last_name), 2)
			p = d.add_paragraph(target.email)

			p = d.add_paragraph()
			# Save a spot to record the email sent date and time in the report
			email_sent_run = p.add_run()

			# Create the Email Opened/Previewed table
			p = d.add_paragraph()
			p.style = d.styles['Normal']
			run = p.add_run("Email Previews")
			run.bold = True

			opened_table = d.add_table(rows=1, cols=1, style="GoReport")
			opened_table.autofit = True
			opened_table.allow_autofit = True

			header1 = opened_table.cell(0,0)
			header1.text = ""
			header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

			# Create the Clicked Link table
			p = d.add_paragraph()
			p.style = d.styles['Normal']
			run = p.add_run("Email Link Clicked")
			run.bold = True

			clicked_table = d.add_table(rows=1, cols=5, style="GoReport")
			clicked_table.autofit = True
			clicked_table.allow_autofit = True

			header1 = clicked_table.cell(0,0)
			header1.text = ""
			header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

			header2 = clicked_table.cell(0,1)
			header2.text = ""
			header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

			header3 = clicked_table.cell(0,2)
			header3.text = ""
			header3.paragraphs[0].add_run("City", "Cell Text").bold = True

			header4 = clicked_table.cell(0,3)
			header4.text = ""
			header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

			header5 = clicked_table.cell(0,4)
			header5.text = ""
			header5.paragraphs[0].add_run("Operating System", "Cell Text").bold = True

			# Create the Submitted Data table
			p = d.add_paragraph()
			p.style = d.styles['Normal']
			run = p.add_run("Phishgate Data Captured")
			run.bold = True

			submitted_table = d.add_table(rows=1, cols=6, style="GoReport")
			submitted_table.autofit = True
			submitted_table.allow_autofit = True

			header1 = submitted_table.cell(0,0)
			header1.text = ""
			header1.paragraphs[0].add_run("Time", "Cell Text").bold = True

			header2 = submitted_table.cell(0,1)
			header2.text = ""
			header2.paragraphs[0].add_run("IP", "Cell Text").bold = True

			header3 = submitted_table.cell(0,2)
			header3.text = ""
			header3.paragraphs[0].add_run("City", "Cell Text").bold = True

			header4 = submitted_table.cell(0,3)
			header4.text = ""
			header4.paragraphs[0].add_run("Browser", "Cell Text").bold = True

			header5 = submitted_table.cell(0,4)
			header5.text = ""
			header5.paragraphs[0].add_run("Operating System", "Cell Text").bold = True

			header6 = submitted_table.cell(0,5)
			header6.text = ""
			header6.paragraphs[0].add_run("Data Captured", "Cell Text").bold = True
			# Parse each timeline event
			# Timestamps are parsed to get date and times by splitting date
			# and time and dropping the milliseconds and timezone
			# Ex: 2017-01-30T14:31:22.534880731-05:00
			for event in self.timeline:
				if event.message == "Email Sent" and event.email == target.email:
					# Parse the timestamp into separate date and time variables
					temp = event.time.split('T')
					sent_date = temp[0]
					sent_time = temp[1].split('.')[0]
					# Record the email sent date and time in the report, in the run reserved earlier
					email_sent_run.text = "Sent on {} at {}".format(sent_date, sent_time)

				if event.message == "Email Opened" and event.email == target.email:
					# Always begin by adding a row to the appropriate table
					opened_table.add_row()
					# Parse the timestamp for and add it to column 0
					# Target the cell located at (counter, 0)
					timestamp = opened_table.cell(opened_counter,0)
					# Get the value for the table cell
					temp = event.time.split('T')
					# Write the value to the table cell
					timestamp.text = temp[0] + " " + temp[1].split('.')[0]
					# Finally, increment the counter to track the row for adding new rows
					# for any addiitonal event sof this type
					opened_counter += 1

				if event.message == "Clicked Link" and event.email == target.email:
					clicked_table.add_row()
					timestamp = clicked_table.cell(clicked_counter,0)
					temp = event.time.split('T')
					timestamp.text = temp[0] + " " + temp[1].split('.')[0]

					ip_add = clicked_table.cell(clicked_counter,1)
					ip_add.text = self.compare_ip_addresses(target.ip, event.details['browser']['address'])

					event_location = clicked_table.cell(clicked_counter,2)
					# Get the location data and add to results row
					# This is based on the IP address pulled from the browser for this event
					# Start by getting the coordinates from GeoLite2
					mmdb_location = self.lookup_ip(event.details['browser']['address'])
					if not mmdb_location == None:
						mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'], mmdb_location['location']['longitude']
						# Check if GoPhish's coordinates agree with these MMDB results
						event_location.text = "{}".format(self.compare_ip_coordinates(target.latitude, target.longitude, mmdb_latitude, mmdb_longitude, event.details['browser']['address']))
					else:
						print("[!] MMDB lookup returned no location results!")
						event_location.text = "IP address look-up returned None"

					# Parse the user-agent string and add browser and OS details to the results row
					user_agent = parse(event.details['browser']['user-agent'])

					browser = clicked_table.cell(clicked_counter, 3)
					browser_details = user_agent.browser.family + " " + user_agent.browser.version_string
					browser.text = browser_details
					self.browsers.append(browser_details)

					op_sys = clicked_table.cell(clicked_counter, 4)
					os_details = user_agent.os.family + " " + user_agent.os.version_string
					op_sys.text = os_details
					self.operating_systems.append(os_details)

					clicked_counter += 1

				if event.message == "Submitted Data" and event.email == target.email:
					submitted_table.add_row()
					timestamp = submitted_table.cell(submitted_counter, 0)
					temp = event.time.split('T')
					timestamp.text = temp[0] + " " + temp[1].split('.')[0]

					ip_add = submitted_table.cell(submitted_counter, 1)
					ip_add.text = event.details['browser']['address']

					event_location = submitted_table.cell(submitted_counter, 2)
					mmdb_location = self.lookup_ip(event.details['browser']['address'])
					if not mmdb_location == None:
						mmdb_latitude, mmdb_longitude = mmdb_location['location']['latitude'], mmdb_location['location']['longitude']
						# Check if GoPhish's coordinates agree with these MMDB results
						event_location.text = "{}".format(self.compare_ip_coordinates(target.latitude, target.longitude, mmdb_latitude, mmdb_longitude, event.details['browser']['address']))
					else:
						print("[!] MMDB lookup returned no location results!")
						event_location.text = "IP address look-up returned None"

					# Parse the user-agent string and add browser and OS details to the results row
					user_agent = parse(event.details['browser']['user-agent'])

					browser = submitted_table.cell(submitted_counter, 3)
					browser_details = user_agent.browser.family + " " + user_agent.browser.version_string
					browser.text = browser_details

					op_sys = submitted_table.cell(submitted_counter, 4)
					os_details = user_agent.os.family + " " + user_agent.os.version_string
					op_sys.text = "{}".format(os_details)

					# Get just the submitted data from the event's payload
					submitted_data = ""
					data = submitted_table.cell(submitted_counter, 5)
					data_payload = event.details['payload']
					# Get all of the submitted data
					for key, value in data_payload.items():
						# To get just submitted data, we drop the 'rid' key
						if not key == "rid":
							submitted_data += "{}:{}   ".format(key, str(value).strip("[").strip("]"))

					data.text = "{}".format(submitted_data)

					submitted_counter += 1

		d.add_page_break()

		# End of the detailed results and the beginning of browser, location, and OS stats
		d.add_heading("Statistics", 1)
		p = d.add_paragraph("The following table shows the browsers seen:")
		# Create browser table
		browser_table = d.add_table(rows=1, cols=2, style="GoReport")
		set_column_width(browser_table.columns[0], Cm(7.24))
		set_column_width(browser_table.columns[1], Cm(3.35))

		header1 = browser_table.cell(0,0)
		header1.text = ""
		header1.paragraphs[0].add_run("Browser", "Cell Text").bold = True

		header2 = browser_table.cell(0,1)
		header2.text =""
		header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

		p = d.add_paragraph("\nThe following table shows the operating systems seen:")

		# Create OS table
		os_table = d.add_table(rows=1, cols=2, style="GoReport")
		set_column_width(browser_table.columns[0], Cm(7.24))
		set_column_width(browser_table.columns[1], Cm(3.35))

		header1 = os_table.cell(0,0)
		header1.text = ""
		header1.paragraphs[0].add_run("Operating System", "Cell Text").bold = True

		header2 = os_table.cell(0,1)
		header2.text =""
		header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

		p = d.add_paragraph("\nThe following table shows the locations seen:")

		# Create geo IP table
		location_table = d.add_table(rows=1, cols=2, style="GoReport")
		set_column_width(browser_table.columns[0], Cm(7.24))
		set_column_width(browser_table.columns[1], Cm(3.35))

		header1 = location_table.cell(0,0)
		header1.text = ""
		header1.paragraphs[0].add_run("Location", "Cell Text").bold = True

		header2 = location_table.cell(0,1)
		header2.text =""
		header2.paragraphs[0].add_run("Visits", "Cell Text").bold = True

		p = d.add_paragraph("\nThe following table shows the IP addresses captured:")

		# Create IP address table
		ip_add_table = d.add_table(rows=1, cols=2, style="GoReport")
		set_column_width(browser_table.columns[0], Cm(7.24))
		set_column_width(browser_table.columns[1], Cm(3.35))

		header1 = ip_add_table.cell(0,0)
		header1.text = ""
		header1.paragraphs[0].add_run("IP Address", "Cell Text").bold = True

		header2 = ip_add_table.cell(0,1)
		header2.text =""
		header2.paragraphs[0].add_run("Seen", "Cell Text").bold = True

		# Counters are used here again to track rows
		counter = 1
		# Counter is used to count all elements in the lists to create a unique list with totals
		counted_browsers = Counter(self.browsers)
		for key, value in counted_browsers.items():
			browser_table.add_row()
			cell = browser_table.cell(counter, 0)
			cell.text = "{}".format(key)

			cell = browser_table.cell(counter, 1)
			cell.text = "{}".format(value)
			counter += 1

		counter = 1
		counted_os = Counter(self.operating_systems)
		for key, value in counted_os.items():
			os_table.add_row()
			cell = os_table.cell(counter, 0)
			cell.text = "{}".format(key)

			cell = os_table.cell(counter, 1)
			cell.text = "{}".format(value)
			counter += 1

		counter = 1
		counted_locations = Counter(self.locations)
		for key, value in counted_locations.items():
			location_table.add_row()
			cell = location_table.cell(counter, 0)
			cell.text = "{}".format(key)

			cell = location_table.cell(counter, 1)
			cell.text = "{}".format(value)
			counter += 1

		counter = 1
		counted_ip_addresses = Counter(self.ip_addresses)
		for key, value in counted_ip_addresses.items():
			ip_add_table.add_row()
			cell = ip_add_table.cell(counter, 0)
			cell.text = "{}".format(key)

			cell = ip_add_table.cell(counter, 1)
			cell.text = "{}".format(value)
			counter += 1

		# Finalize document and save it as the value of output_word_report
		d.save("{}".format(self.output_word_report))
		print("[+] Done! Check \"{}\" for your results.".format(self.output_word_report))



if __name__ == '__main__':
	parse_options()
